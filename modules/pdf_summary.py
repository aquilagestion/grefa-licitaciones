"""Resumen / comprobación de pliegos y ofertas con Gemini (PDF, Word, Excel)."""

from __future__ import annotations

import logging
import os
import re
import time
from pathlib import Path
from typing import Any

from modules.sheets_store import _secret

LOGGER = logging.getLogger(__name__)

# Preferir Flash «clásico»: gemini-3.x free tier suele ser solo ~20 req/día.
DEFAULT_MODEL = "gemini-2.5-flash"
DEFAULT_FALLBACK_MODELS = (
    "gemini-flash-latest",
    "gemini-2.5-flash-lite",
    "gemini-2.0-flash",
)
MAX_PDF_BYTES = 15 * 1024 * 1024  # ~15 MB
MAX_TEXTO_EXTRAIDO = 120_000
MAX_REINTENTOS_429 = 1
MAX_ESPERA_429_S = 45.0

#: Extensiones aceptadas en uploaders y validación.
EXTENSIONES_DOC = ("pdf", "docx", "xlsx")
MIME_POR_EXT = {
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
}
PROMPT_PLIEGO = """Eres un analista de licitaciones públicas para el equipo GREFA
(Grupo para la Recuperación de la Fauna Autóctona).

Analiza el/los pliego(s) adjunto(s). Puede haber varios documentos (p. ej. PCAP =
Pliego de Cláusulas Administrativas Particulares y PPT = Pliego de Prescripciones
Técnicas). Intégralos en un único resumen en **español**, claro y accionable,
con estas secciones (usa títulos markdown):

## Objeto del contrato
## Presupuesto e importes (si constan)
## Requisitos de solvencia y clasificación (sobre todo PCAP)
## Criterios de adjudicación
## Alcance técnico / prestaciones (sobre todo PPT)
## Plazos clave (presentación, ejecución, garantías)
## Documentación a presentar
## Puntos de atención para GREFA
## Recomendación breve (Presentar / Estudiar / Descartar) y por qué

Si algún dato no aparece en los documentos, indícalo como «No consta».
Sé conciso pero no omitas requisitos eliminatorios. Indica de qué documento
sale cada dato clave cuando haya varios."""

PROMPT_TEXTO = """Eres un analista de licitaciones públicas para GREFA.
Resume el siguiente extracto de pliego(s) en español con las mismas secciones que
un pliego completo (objeto, solvencia, criterios, alcance técnico, plazos,
documentación, atención GREFA, recomendación). Si falta información, indica «No consta».

---
{texto}
"""

# Límites de envío a Gemini en el comprobador (parcial por diseño).
MAX_OFERTA_COMPROBADOR = 4
MAX_PLIEGO_COMPROBADOR = 3

PROMPT_COMPROBADOR = """Eres un revisor de documentación de oferta para licitaciones públicas
españolas (PLACSP), trabajando para GREFA.

Te pasan:
1) Uno o varios documentos de **oferta** que GREFA está preparando (puede ser un
   subconjunto: p. ej. 3–4 de 16). NO asumas que es el paquete completo.
2) Opcionalmente, documentos de referencia del expediente, en especial:
   - **Cláusulas / condiciones administrativas** (PCAP, PCP, condiciones particulares…)
   - **Prescripciones técnicas** (PPT / pliego de prescripciones técnicas)
   - Anejos y, si existe, ficha resumen PLACSP («Pliegos.pdf» de la plataforma)
3) Opcionalmente, checklist / requisitos adicionales del usuario.

## Objetivo (muy importante)
Comprobar si **los documentos subidos se adaptan / conforman** a lo exigido por:
- las **prescripciones / cláusulas administrativas** (formas, sobres, DECLAREs,
  solvencia, modelos administrativos, firmas, sobres electrónicos, etc.), y
- las **prescripciones técnicas** (memoria, metodología, medios, criterios
  técnicos, anexos técnicos, etc.),
para **el tipo de documento** revisado (contenido, estructura, campos, criterios,
referencias al expediente).

NO castigues el veredicto porque falten otros documentos del paquete que el
usuario **no ha subido**. Eso va solo a la sección de cobertura del paquete.

Si solo hay ficha resumen PLACSP (objeto, CPV, importes, plazos, enlaces), úsala
como contexto; si faltan PCAP y/o PPT completos, dilo en limitaciones y no
inventes modelos.

No inventes datos que no estén en los ficheros.
Si no hay referencia administrativa ni técnica, evalúa solo formalidad interna
y marca limitaciones.

Responde en **español**, con markdown y estas secciones exactas:

## Veredicto
Una sola línea al inicio con UNA de estas etiquetas (sobre **lo subido**, no
sobre el paquete completo):
- `✅ CONFORME` — lo revisado encaja con administrativas y/o técnicas aplicables
  sin defectos bloqueantes evidentes
- `⚠️ CONFORME CON RESERVAS` — útiles pero con huecos, dudas o riesgos en lo subido
- `❌ NO CONFORME` — errores graves o clara desadaptación a PCAP/PPT **en
  los documentos aportados**

Prohibido usar «NO APTO / INCOMPLETO» solo porque el paquete esté incompleto.
Luego 2–4 frases justificando el veredicto (centradas en conformidad admin/técnica).

## Errores y defectos en lo subido
Lista numerada. Cada ítem: gravedad (`Bloqueante` / `Importante` / `Menor`),
si proviene de **administrativas** o **técnicas**, qué está mal y en qué
documento/sección. Si no hay, «Ninguno evidente».

## Adaptación a cláusulas administrativas
Para cada documento de oferta que corresponda a exigencias administrativas:
a qué cláusula/modelo del PCAP/PCP parece responder y si encaja (sí / parcial / no).
Si no aplica a ninguno, indícalo.

## Adaptación a prescripciones técnicas
Igual respecto al PPT / prescripciones técnicas (memoria, criterios, medios, etc.).
Si no aplica a ninguno, indícalo.

## Cobertura del paquete (informativo; no decide el veredicto)
Lista de documentación que PCAP/PPT/checklist exigen y que **no** está entre
los PDF de oferta subidos. Orientación para completar el expediente, no un fallo
de los ficheros revisados.

## Inconsistencias y riesgos
Fechas, importes, NIF, firmas, sobres, referencias cruzadas, formatos, etc.
en lo aportado.

## Checklist sobre los documentos revisados
5–10 casillas accionables (`[ ] …`) para mejorar **estos** ficheros.

## Limitaciones de este análisis
Qué no has podido verificar (falta PCAP o PPT, escaneos, firmas, DEH, etc.).

Sé concreto. El veredicto mide conformidad de lo subido frente a administrativas
y técnicas, no la completitud del lote."""


class PdfSummaryError(RuntimeError):
    """Error al procesar o resumir un documento (PDF / Word / Excel)."""


def api_key() -> str | None:
    clave = _secret("gemini", "api_key") or os.environ.get("GREFA_GEMINI_API_KEY")
    return str(clave).strip() if clave else None


def model_name() -> str:
    return str(
        _secret("gemini", "model")
        or os.environ.get("GREFA_GEMINI_MODEL")
        or DEFAULT_MODEL
    ).strip()


def model_fallbacks() -> list[str]:
    """Modelos alternativos si el principal agota cuota (429)."""
    crudo = _secret("gemini", "fallback_models") or os.environ.get(
        "GREFA_GEMINI_FALLBACK_MODELS"
    )
    if isinstance(crudo, (list, tuple)):
        extras = [str(x).strip() for x in crudo if str(x).strip()]
    elif crudo:
        extras = [x.strip() for x in str(crudo).replace(";", ",").split(",") if x.strip()]
    else:
        extras = list(DEFAULT_FALLBACK_MODELS)
    principal = model_name()
    vistos = {principal}
    salida: list[str] = []
    for m in extras:
        if m not in vistos:
            vistos.add(m)
            salida.append(m)
    return salida


def groq_api_key() -> str | None:
    clave = _secret("groq", "api_key") or os.environ.get("GREFA_GROQ_API_KEY")
    return str(clave).strip() if clave else None


def groq_model() -> str:
    return str(
        _secret("groq", "model")
        or os.environ.get("GREFA_GROQ_MODEL")
        or "llama-3.3-70b-versatile"
    ).strip()


def openrouter_api_key() -> str | None:
    clave = (
        _secret("openrouter", "api_key")
        or os.environ.get("GREFA_OPENROUTER_API_KEY")
    )
    return str(clave).strip() if clave else None


def openrouter_model() -> str:
    return str(
        _secret("openrouter", "model")
        or os.environ.get("GREFA_OPENROUTER_MODEL")
        or "meta-llama/llama-3.3-70b-instruct:free"
    ).strip()


def is_configured() -> bool:
    """True si hay al menos un proveedor de IA gratuito configurado."""
    return bool(api_key() or groq_api_key() or openrouter_api_key())


def proveedores_configurados() -> list[str]:
    nombres: list[str] = []
    if api_key():
        nombres.append(f"Gemini ({model_name()})")
    if groq_api_key():
        nombres.append(f"Groq ({groq_model()})")
    if openrouter_api_key():
        nombres.append(f"OpenRouter ({openrouter_model()})")
    return nombres


# Aviso emergente al cambiar de proveedor (toast + banner en la UI).
_ULTIMO_AVISO_PROVEEDOR: str | None = None


def consumir_aviso_proveedor() -> str | None:
    global _ULTIMO_AVISO_PROVEEDOR
    msg = _ULTIMO_AVISO_PROVEEDOR
    _ULTIMO_AVISO_PROVEEDOR = None
    return msg


def _avisar_cambio_proveedor(mensaje: str) -> None:
    """Toast emergente + aviso persistente en session_state."""
    global _ULTIMO_AVISO_PROVEEDOR
    _ULTIMO_AVISO_PROVEEDOR = mensaje
    LOGGER.warning("%s", mensaje)
    try:
        import streamlit as st

        st.session_state["ia_aviso_proveedor"] = mensaje
        st.toast(mensaje, icon="⚠️")
    except Exception:
        pass


def mostrar_avisos_ia() -> None:
    """Muestra en pantalla el aviso de cambio de proveedor si lo hay."""
    try:
        import streamlit as st
    except ImportError:
        return
    msg = st.session_state.pop("ia_aviso_proveedor", None) or consumir_aviso_proveedor()
    if msg:
        st.warning(msg)


def _es_error_cuota(exc: BaseException) -> bool:
    texto = str(exc).lower()
    return (
        "429" in texto
        or "resource_exhausted" in texto
        or "quota" in texto
        or "rate limit" in texto
        or "rate-limit" in texto
    )


def _segundos_reintento(exc: BaseException) -> float | None:
    texto = str(exc)
    m = re.search(r"retry in\s+(\d+(?:\.\d+)?)\s*s", texto, flags=re.I)
    if m:
        return float(m.group(1))
    m = re.search(r"retry_delay\s*\{\s*seconds:\s*(\d+)", texto, flags=re.I)
    if m:
        return float(m.group(1))
    return None


def _es_cuota_diaria(exc: BaseException) -> bool:
    texto = str(exc).lower().replace("_", "")
    return "perday" in texto or "requestsperday" in texto


def _mensaje_sin_proveedores(exc: BaseException | None = None) -> str:
    extras = ""
    if exc:
        extras = f" Último error: {exc}"
    return (
        "Ningún proveedor de IA disponible (cuota agotada o no configurado). "
        "Configura en Secrets al menos uno: [gemini], [groq] o [openrouter] "
        "(tier gratuito). Orden de uso: Gemini → Groq → OpenRouter."
        + extras
    )


def _extension(nombre: str) -> str:
    return Path(nombre or "").suffix.lower().lstrip(".")


def _extraer_texto_pdf(pdf_bytes: bytes) -> str:
    try:
        from pypdf import PdfReader
        from io import BytesIO

        lector = PdfReader(BytesIO(pdf_bytes))
        partes: list[str] = []
        for pagina in lector.pages:
            texto = pagina.extract_text() or ""
            if texto.strip():
                partes.append(texto.strip())
            if sum(len(p) for p in partes) >= MAX_TEXTO_EXTRAIDO:
                break
        return "\n\n".join(partes)[:MAX_TEXTO_EXTRAIDO]
    except Exception as exc:
        raise PdfSummaryError(f"No se pudo leer el PDF: {exc}") from exc


def _extraer_texto_docx(datos: bytes) -> str:
    try:
        from io import BytesIO

        from docx import Document
    except ImportError as exc:
        raise PdfSummaryError("Falta python-docx para leer Word (.docx).") from exc
    try:
        doc = Document(BytesIO(datos))
        partes: list[str] = []
        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if t:
                partes.append(t)
            if sum(len(x) for x in partes) >= MAX_TEXTO_EXTRAIDO:
                break
        for tabla in doc.tables:
            for fila in tabla.rows:
                celdas = [((c.text or "").strip()) for c in fila.cells]
                if any(celdas):
                    partes.append(" | ".join(celdas))
            if sum(len(x) for x in partes) >= MAX_TEXTO_EXTRAIDO:
                break
        return "\n".join(partes)[:MAX_TEXTO_EXTRAIDO]
    except Exception as exc:
        raise PdfSummaryError(f"No se pudo leer el Word (.docx): {exc}") from exc


def _extraer_texto_xlsx(datos: bytes) -> str:
    try:
        from io import BytesIO

        from openpyxl import load_workbook
    except ImportError as exc:
        raise PdfSummaryError("Falta openpyxl para leer Excel (.xlsx).") from exc
    try:
        wb = load_workbook(BytesIO(datos), read_only=True, data_only=True)
        partes: list[str] = []
        for hoja in wb.worksheets:
            partes.append(f"## Hoja: {hoja.title}")
            for fila in hoja.iter_rows(values_only=True):
                vals = ["" if v is None else str(v).strip() for v in fila]
                if any(vals):
                    partes.append(" | ".join(vals))
                if sum(len(x) for x in partes) >= MAX_TEXTO_EXTRAIDO:
                    break
            if sum(len(x) for x in partes) >= MAX_TEXTO_EXTRAIDO:
                break
        try:
            wb.close()
        except Exception:
            pass
        return "\n".join(partes)[:MAX_TEXTO_EXTRAIDO]
    except Exception as exc:
        raise PdfSummaryError(f"No se pudo leer el Excel (.xlsx): {exc}") from exc


def _extraer_texto_documento(doc: dict[str, Any]) -> str:
    """Extrae texto según extensión (pdf / docx / xlsx)."""
    nombre = str(doc.get("nombre") or "documento.pdf")
    datos = doc.get("bytes") or b""
    ext = _extension(nombre)
    if ext == "pdf":
        return _extraer_texto_pdf(datos)
    if ext == "docx":
        return _extraer_texto_docx(datos)
    if ext == "xlsx":
        return _extraer_texto_xlsx(datos)
    if ext in {"doc", "xls"}:
        raise PdfSummaryError(
            f"{nombre}: el formato .{ext} antiguo no está soportado. "
            "Guárdalo como .docx o .xlsx e inténtalo de nuevo."
        )
    raise PdfSummaryError(
        f"{nombre}: extensión no soportada. Usa PDF, Word (.docx) o Excel (.xlsx)."
    )


def _contenido_a_texto(contenido: list[Any]) -> str:
    """Convierte partes multimodal (PDF Gemini) a texto para Groq/OpenRouter."""
    bloques: list[str] = []
    for item in contenido or []:
        if isinstance(item, str):
            if item.strip():
                bloques.append(item.strip())
            continue
        if isinstance(item, dict) and item.get("data"):
            mime = str(item.get("mime_type") or "")
            datos = item.get("data") or b""
            if "pdf" in mime and isinstance(datos, (bytes, bytearray)):
                try:
                    texto = _extraer_texto_pdf(bytes(datos))
                    if texto.strip():
                        bloques.append(texto.strip())
                    else:
                        bloques.append("[PDF sin texto extraíble]")
                except Exception as exc:
                    bloques.append(f"[PDF no legible: {exc}]")
            else:
                bloques.append(f"[Adjunto {mime or 'binario'} omitido]")
    return "\n\n".join(bloques)[:MAX_TEXTO_EXTRAIDO]


def _generar_openai_compatible(
    *,
    proveedor: str,
    base_url: str,
    api_key: str,
    model: str,
    prompt: str,
    headers_extra: dict[str, str] | None = None,
) -> str:
    """Chat completions (Groq / OpenRouter). Solo texto."""
    import requests

    url = base_url.rstrip("/") + "/chat/completions"
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
    }
    if headers_extra:
        headers.update(headers_extra)
    cuerpo = {
        "model": model,
        "messages": [
            {
                "role": "system",
                "content": (
                    "Eres analista de contratación pública española para GREFA. "
                    "Responde siempre en español, con markdown claro y accionable."
                ),
            },
            {"role": "user", "content": prompt[:MAX_TEXTO_EXTRAIDO]},
        ],
        "temperature": 0.2,
    }
    try:
        resp = requests.post(url, headers=headers, json=cuerpo, timeout=180)
    except requests.RequestException as exc:
        raise PdfSummaryError(f"Error de red con {proveedor}: {exc}") from exc

    if resp.status_code == 429 or (
        resp.status_code == 402
    ):  # OpenRouter sin créditos / rate limit
        raise PdfSummaryError(
            f"429 cuota {proveedor}: {resp.text[:500]}"
        )
    if resp.status_code >= 400:
        raise PdfSummaryError(
            f"Error {proveedor} HTTP {resp.status_code}: {resp.text[:800]}"
        )
    try:
        data = resp.json()
        texto = (
            data.get("choices", [{}])[0]
            .get("message", {})
            .get("content", "")
        )
    except Exception as exc:
        raise PdfSummaryError(f"Respuesta inválida de {proveedor}: {exc}") from exc
    texto = (texto or "").strip()
    if not texto:
        raise PdfSummaryError(f"{proveedor} no devolvió texto.")
    return texto


def _intentar_gemini(
    contenido: list[Any],
    *,
    prompt: str,
) -> str:
    clave = api_key()
    if not clave:
        raise PdfSummaryError("Gemini no configurado.")

    try:
        import google.generativeai as genai
    except ImportError as exc:
        raise PdfSummaryError("Falta el paquete google-generativeai.") from exc

    genai.configure(api_key=clave)
    candidatos = [model_name(), *model_fallbacks()]
    ultimo_exc: BaseException | None = None
    ultimo_modelo = candidatos[0]

    for idx, nombre_modelo in enumerate(candidatos):
        ultimo_modelo = nombre_modelo
        modelo = genai.GenerativeModel(nombre_modelo)
        for intento in range(MAX_REINTENTOS_429 + 1):
            try:
                respuesta = modelo.generate_content([*contenido, prompt])
                texto = (respuesta.text or "").strip()
                if not texto:
                    raise PdfSummaryError(
                        "Gemini no devolvió texto. Prueba con otro PDF o modelo."
                    )
                if idx > 0:
                    _avisar_cambio_proveedor(
                        f"Gemini principal sin cuota. "
                        f"Se está usando el modelo Gemini «{nombre_modelo}»."
                    )
                return texto
            except PdfSummaryError:
                raise
            except Exception as exc:
                ultimo_exc = exc
                if not _es_error_cuota(exc):
                    raise PdfSummaryError(f"Error al llamar a Gemini: {exc}") from exc
                espera = _segundos_reintento(exc)
                if (
                    intento < MAX_REINTENTOS_429
                    and espera is not None
                    and espera <= MAX_ESPERA_429_S
                    and not _es_cuota_diaria(exc)
                ):
                    time.sleep(min(espera + 1.0, MAX_ESPERA_429_S))
                    continue
                break

    assert ultimo_exc is not None
    raise PdfSummaryError(
        f"Cuota Gemini agotada ({ultimo_modelo}): {ultimo_exc}"
    ) from ultimo_exc


def _generar_con_gemini(
    contenido: list[Any],
    *,
    contexto: str = "",
    prompt_base: str | None = None,
) -> str:
    """Genera texto con cadena gratuita: Gemini → Groq → OpenRouter."""
    if not is_configured():
        raise PdfSummaryError(
            "Ninguna IA configurada. Añade en Secrets [gemini] api_key y/o "
            "[groq] api_key y/o [openrouter] api_key (tier gratuito)."
        )

    base = prompt_base or PROMPT_PLIEGO
    prompt = base
    if contexto:
        prompt = f"{base}\n\nContexto del expediente:\n{contexto}"

    errores: list[str] = []
    gemini_intento = False

    # 1) Gemini (PDF nativo + modelos fallback internos)
    if api_key():
        gemini_intento = True
        try:
            return _intentar_gemini(contenido, prompt=prompt)
        except PdfSummaryError as exc:
            errores.append(str(exc))
            if not _es_error_cuota(exc) and "no configurado" not in str(exc).lower():
                # Error no-cuota: aún así probar otros si hay; si no, relanzar
                if not (groq_api_key() or openrouter_api_key()):
                    raise
            LOGGER.warning("Gemini no disponible; se prueba Groq/OpenRouter. %s", exc)

    texto_docs = _contenido_a_texto(contenido)
    prompt_texto = prompt
    if texto_docs.strip():
        prompt_texto = (
            f"{prompt}\n\n---\nDocumentos (texto extraído):\n{texto_docs}\n---"
        )

    # 2) Groq
    if groq_api_key():
        if gemini_intento:
            _avisar_cambio_proveedor(
                "Cuota o fallo de Gemini. Se usará Groq (Llama, tier gratuito) "
                "para esta petición."
            )
        else:
            _avisar_cambio_proveedor(
                "Gemini no está configurado. Se usará Groq (Llama, tier gratuito)."
            )
        try:
            return _generar_openai_compatible(
                proveedor="Groq",
                base_url="https://api.groq.com/openai/v1",
                api_key=groq_api_key() or "",
                model=groq_model(),
                prompt=prompt_texto,
            )
        except PdfSummaryError as exc:
            errores.append(str(exc))
            if not _es_error_cuota(exc) and not openrouter_api_key():
                raise
            LOGGER.warning("Groq no disponible; se prueba OpenRouter. %s", exc)

    # 3) OpenRouter (modelos :free)
    if openrouter_api_key():
        if gemini_intento or groq_api_key():
            _avisar_cambio_proveedor(
                "Gemini/Groq sin cuota o no disponibles. "
                "Se usará OpenRouter (modelo gratuito) para esta petición."
            )
        else:
            _avisar_cambio_proveedor(
                "Solo OpenRouter configurado. Se usará su modelo gratuito."
            )
        try:
            return _generar_openai_compatible(
                proveedor="OpenRouter",
                base_url="https://openrouter.ai/api/v1",
                api_key=openrouter_api_key() or "",
                model=openrouter_model(),
                prompt=prompt_texto,
                headers_extra={
                    "HTTP-Referer": "https://huggingface.co/spaces",
                    "X-Title": "GREFA Licitaciones",
                },
            )
        except PdfSummaryError as exc:
            errores.append(str(exc))
            raise PdfSummaryError(_mensaje_sin_proveedores(exc)) from exc

    raise PdfSummaryError(
        _mensaje_sin_proveedores(Exception(" | ".join(errores) if errores else None))
    )


def _validar_pdfs(documentos: list[dict[str, Any]], *, etiqueta: str) -> list[dict[str, Any]]:
    """Valida PDF / Word / Excel (nombre histórico por compatibilidad)."""
    utiles = [
        d
        for d in documentos
        if isinstance(d, dict) and d.get("bytes") and len(d.get("bytes") or b"") > 0
    ]
    if not utiles:
        raise PdfSummaryError(
            f"No hay documentos válidos en {etiqueta} (PDF, DOCX o XLSX)."
        )
    for doc in utiles:
        nombre = str(doc.get("nombre") or "documento")
        ext = _extension(nombre)
        if ext not in EXTENSIONES_DOC and ext not in {"doc", "xls"}:
            raise PdfSummaryError(
                f"{nombre}: formato no soportado. Usa PDF, Word (.docx) o Excel (.xlsx)."
            )
        if len(doc["bytes"]) > MAX_PDF_BYTES:
            raise PdfSummaryError(
                f"{nombre} supera {MAX_PDF_BYTES // (1024 * 1024)} MB."
            )
    return utiles


def _partes_gemini_pdfs(documentos: list[dict[str, Any]], *, max_docs: int = 4) -> list[Any]:
    """Prepara partes para Gemini: PDF nativo; Word/Excel como texto extraído."""
    partes: list[Any] = []
    for doc in documentos[:max_docs]:
        nombre = str(doc.get("nombre") or "documento.pdf")
        etiqueta = f"[{doc.get('tipo', 'OTRO')}] {nombre}"
        partes.append(etiqueta)
        ext = _extension(nombre)
        if ext == "pdf":
            partes.append({"mime_type": "application/pdf", "data": doc["bytes"]})
        else:
            # Word/Excel: texto (Gemini no siempre acepta estos MIME con fiabilidad).
            texto = _extraer_texto_documento(doc)
            if len(texto.strip()) < 20:
                raise PdfSummaryError(
                    f"{nombre} no tiene texto extraíble (vacío o solo imágenes)."
                )
            partes.append(
                f"Contenido extraído de {nombre} ({ext}):\n---\n{texto}\n---"
            )
    return partes


def _texto_desde_pdfs(documentos: list[dict[str, Any]]) -> str:
    bloques: list[str] = []
    for doc in documentos:
        try:
            texto = _extraer_texto_documento(doc)
        except PdfSummaryError:
            continue
        if texto.strip():
            bloques.append(
                f"### {doc.get('tipo', 'OTRO')} — {doc.get('nombre', 'documento')}\n{texto}"
            )
    return "\n\n".join(bloques)[:MAX_TEXTO_EXTRAIDO]


def summarize_pdf(
    pdf_bytes: bytes,
    *,
    expediente: str = "",
    titulo: str = "",
) -> str:
    """Resume un pliego PDF. Intenta enviar el PDF a Gemini; si falla, extrae texto."""
    return summarize_documentos(
        [{"nombre": "pliego.pdf", "tipo": "OTRO", "bytes": pdf_bytes}],
        expediente=expediente,
        titulo=titulo,
    )


def summarize_documentos(
    documentos: list[dict[str, Any]],
    *,
    expediente: str = "",
    titulo: str = "",
) -> str:
    """Resume uno o varios PDFs (p. ej. PCAP + PPT) en un único informe."""
    utiles = _validar_pdfs(documentos, etiqueta="pliegos")

    contexto_parts = []
    if expediente or titulo:
        contexto_parts.append(f"Expediente: {expediente or '—'}\nTítulo: {titulo or '—'}")
    listado = ", ".join(
        f"{d.get('tipo', 'OTRO')}: {d.get('nombre', 'documento.pdf')}" for d in utiles
    )
    contexto_parts.append(f"Documentos analizados: {listado}")
    contexto = "\n".join(contexto_parts)

    try:
        return _generar_con_gemini(
            _partes_gemini_pdfs(utiles),
            contexto=contexto,
            prompt_base=PROMPT_PLIEGO,
        )
    except PdfSummaryError as exc:
        LOGGER.info("PDF multi directo falló (%s); se intenta texto.", exc)

    combinado = _texto_desde_pdfs(utiles)
    if len(combinado.strip()) < 200:
        raise PdfSummaryError(
            "No se pudo extraer texto suficiente. "
            "Sube PDF con texto, Word (.docx) o Excel (.xlsx) nativos."
        )

    prompt = PROMPT_TEXTO.format(texto=combinado)
    if contexto:
        prompt = f"{contexto}\n\n{prompt}"
    return _generar_con_gemini([prompt], prompt_base=PROMPT_PLIEGO)


def comprobar_documentos(
    documentos_oferta: list[dict[str, Any]],
    *,
    documentos_pliego: list[dict[str, Any]] | None = None,
    expediente: str = "",
    titulo: str = "",
    requisitos_texto: str = "",
) -> str:
    """Comprueba si los PDF de oferta se adaptan al pliego (revisión parcial)."""
    oferta = _validar_pdfs(documentos_oferta, etiqueta="documentos de oferta")
    pliegos = []
    if documentos_pliego:
        try:
            pliegos = _validar_pdfs(documentos_pliego, etiqueta="pliegos de referencia")
        except PdfSummaryError:
            pliegos = []

    omitidos_oferta = max(0, len(oferta) - MAX_OFERTA_COMPROBADOR)
    omitidos_pliego = max(0, len(pliegos) - MAX_PLIEGO_COMPROBADOR)
    oferta_envio = oferta[:MAX_OFERTA_COMPROBADOR]
    pliegos_envio = pliegos[:MAX_PLIEGO_COMPROBADOR]

    contexto_parts = [
        "Modo: COMPROBADOR DE CONFORMIDAD (revisión parcial).",
        "El usuario puede haber subido solo parte del paquete. "
        "El veredicto mide si LO SUBIDO se adapta a las cláusulas/prescripciones "
        "ADMINISTRATIVAS (PCAP/PCP) y a las PRESCRIPCIONES TÉCNICAS (PPT); "
        "NO marques NO CONFORME solo por documentos no aportados.",
        f"Expediente: {expediente or '—'}",
        f"Título / referencia: {titulo or '—'}",
        "Documentos de OFERTA subidos (revisar conformidad): "
        + ", ".join(
            f"{d.get('tipo', 'OFERTA')}: {d.get('nombre', 'documento.pdf')}"
            for d in oferta_envio
        ),
        f"Total documentos oferta seleccionados: {len(oferta)}"
        + (
            f" (se analizan los {len(oferta_envio)} primeros; {omitidos_oferta} omitidos)"
            if omitidos_oferta
            else ""
        ),
    ]
    if pliegos_envio:
        contexto_parts.append(
            "Referencia del expediente (prioriza PCAP/administrativas y PPT/técnicas): "
            + ", ".join(
                f"{d.get('tipo', 'PLIEGO')}: {d.get('nombre', 'documento.pdf')}"
                for d in pliegos_envio
            )
        )
        if omitidos_pliego:
            contexto_parts.append(
                f"Se omiten {omitidos_pliego} documentos de referencia por límite de análisis."
            )
    else:
        contexto_parts.append(
            "No se adjuntaron cláusulas administrativas ni prescripciones técnicas: "
            "evalúa solo formalidad interna y señala limitaciones."
        )
    if requisitos_texto and requisitos_texto.strip():
        contexto_parts.append(
            "Checklist / requisitos adicionales indicados por el usuario:\n"
            + requisitos_texto.strip()[:8000]
        )
    contexto = "\n".join(contexto_parts)

    docs_envio = [
        {**d, "tipo": d.get("tipo") or "OFERTA"} for d in oferta_envio
    ] + [
        {**d, "tipo": d.get("tipo") or "PLIEGO"} for d in pliegos_envio
    ]
    max_docs = MAX_OFERTA_COMPROBADOR + MAX_PLIEGO_COMPROBADOR

    try:
        partes: list[Any] = ["=== DOCUMENTACIÓN A REVISAR (CONFORMIDAD) ==="]
        partes.extend(_partes_gemini_pdfs(docs_envio, max_docs=max_docs))
        return _generar_con_gemini(
            partes,
            contexto=contexto,
            prompt_base=PROMPT_COMPROBADOR,
        )
    except PdfSummaryError as exc:
        LOGGER.info("Comprobador PDF directo falló (%s); se intenta texto.", exc)

    bloques = ["=== OFERTA ===", _texto_desde_pdfs(oferta_envio)]
    if pliegos_envio:
        bloques.extend(["=== PLIEGO / REQUISITOS ===", _texto_desde_pdfs(pliegos_envio)])
    combinado = "\n\n".join(b for b in bloques if b).strip()[:MAX_TEXTO_EXTRAIDO]
    if len(combinado) < 200:
        raise PdfSummaryError(
            "No se pudo extraer texto suficiente de los documentos. "
            "Sube PDF con texto, Word (.docx) o Excel (.xlsx)."
        )

    return _generar_con_gemini(
        [f"Extracto de documentos:\n---\n{combinado}\n---"],
        contexto=contexto,
        prompt_base=PROMPT_COMPROBADOR,
    )


PROMPT_SINTESIS_LOTES = """Eres un revisor de documentación de oferta para licitaciones públicas
españolas (PLACSP), trabajando para GREFA.

Te pasan **varios informes parciales** de conformidad. Cada uno analiza un lote
de documentos (máx. 4) del **mismo expediente**. NO tienes los PDF originales:
solo los informes parciales.

## Objetivo
Elaborar **un único informe global** que unifique todos los lotes: veredicto
conjunto, errores, adaptación a cláusulas administrativas y prescripciones
técnicas, cobertura del paquete e inconsistencias cruzadas entre lotes.

Reglas de veredicto global:
- Si algún lote es `❌ NO CONFORME` por defectos en lo subido → el global no
  puede ser `✅ CONFORME` (como mínimo `⚠️ CONFORME CON RESERVAS` o `❌ NO CONFORME`).
- No castigues por documentos no aportados en ningún lote (cobertura informativa).
- Si los lotes se complementan bien, dilo explícitamente.

Usa las mismas secciones markdown que un informe de conformidad:

## Veredicto
Una línea con `✅ CONFORME` / `⚠️ CONFORME CON RESERVAS` / `❌ NO CONFORME`,
luego 2–4 frases.

## Errores y defectos en lo subido
## Adaptación a cláusulas administrativas
## Adaptación a prescripciones técnicas
## Cobertura del paquete (informativo; no decide el veredicto)
## Inconsistencias y riesgos
## Checklist sobre los documentos revisados
## Limitaciones de este análisis
(Incluye que el global se basa en síntesis de informes parciales, no en relectura de PDF.)

Sé concreto y accionable. En español."""


def sintetizar_informes_comprobador(
    informes_parciales: list[dict[str, Any]],
    *,
    expediente: str = "",
    titulo: str = "",
    requisitos_texto: str = "",
) -> str:
    """Unifica varios informes parciales de lotes en un informe global."""
    utiles = [
        i
        for i in (informes_parciales or [])
        if isinstance(i, dict) and str(i.get("informe") or "").strip()
    ]
    if not utiles:
        raise PdfSummaryError("No hay informes parciales para sintetizar.")
    if len(utiles) == 1:
        # Un solo lote: el global es ese informe (marcado como tal).
        unico = str(utiles[0]["informe"]).strip()
        cabecera = (
            "_Informe global a partir de 1 lote (sin síntesis adicional)._\n\n"
        )
        return cabecera + unico

    bloques = []
    for i, item in enumerate(utiles, start=1):
        nombres = item.get("nombres") or []
        lista = ", ".join(str(n) for n in nombres) if nombres else "—"
        bloques.append(
            f"### Informe parcial — Lote {i}\n"
            f"Documentos: {lista}\n\n"
            f"{str(item.get('informe') or '').strip()}"
        )
    cuerpo = "\n\n---\n\n".join(bloques)[:MAX_TEXTO_EXTRAIDO]

    contexto_parts = [
        "Modo: SÍNTESIS DE LOTES DEL COMPROBADOR.",
        f"Expediente: {expediente or '—'}",
        f"Título / referencia: {titulo or '—'}",
        f"Número de lotes parciales: {len(utiles)}",
    ]
    if requisitos_texto and requisitos_texto.strip():
        contexto_parts.append(
            "Checklist / requisitos del usuario:\n" + requisitos_texto.strip()[:4000]
        )
    contexto = "\n".join(contexto_parts)

    return _generar_con_gemini(
        [
            "=== INFORMES PARCIALES A UNIFICAR ===\n\n" + cuerpo,
        ],
        contexto=contexto,
        prompt_base=PROMPT_SINTESIS_LOTES,
    )


PROMPT_CONVOCATORIA_AYUDAS = """Eres analista de convocatorias de ayudas, subvenciones y premios
públicos (BDNS / bases reguladoras) para el equipo GREFA
(Grupo para la Recuperación de la Fauna Autóctona).

A partir del texto de la convocatoria (ficha BDNS, bases, extracto web…),
extrae en **español** un informe claro y accionable con estas secciones exactas
(markdown):

## Requisitos de participación / elegibilidad
Lista numerada de quién puede presentarse, condiciones, exclusiones,
ámbito territorial, tipología de beneficiario, etc. Si no consta, «No consta».

## Documentación a entregar
Lista numerada de documentos, formularios, anexos, memorias, certificados
u otros materiales que haya que presentar. Si no consta, «No consta».

## Plazos clave
Presentación, resolución u otros plazos mencionados. Si no consta, «No consta».

## Sede / enlaces útiles
Sede electrónica u otros enlaces citados. Si no consta, «No consta».

## Atención para GREFA
2–5 puntos de riesgo u oportunidad (idoneidad, carga documental, plazos cortos…).

No inventes datos que no estén en el texto. Sé conciso y práctico.
"""


def extraer_requisitos_convocatoria(
    texto: str,
    *,
    expediente: str = "",
    titulo: str = "",
    url: str = "",
) -> str:
    """Extrae requisitos y documentación a entregar de una convocatoria (texto)."""
    cuerpo = (texto or "").strip()
    if len(cuerpo) < 80:
        raise PdfSummaryError(
            "No hay texto suficiente de la convocatoria para extraer requisitos. "
            "Abre la ficha oficial o actualiza datos BDNS."
        )
    contexto_parts = []
    if expediente or titulo:
        contexto_parts.append(
            f"Código / expediente: {expediente or '—'}\nTítulo: {titulo or '—'}"
        )
    if url:
        contexto_parts.append(f"URL: {url}")
    contexto = "\n".join(contexto_parts)
    prompt = (
        f"{PROMPT_CONVOCATORIA_AYUDAS}\n\n---\nTexto de la convocatoria:\n"
        f"{cuerpo[:MAX_TEXTO_EXTRAIDO]}\n---"
    )
    return _generar_con_gemini([prompt], contexto=contexto, prompt_base=PROMPT_CONVOCATORIA_AYUDAS)
