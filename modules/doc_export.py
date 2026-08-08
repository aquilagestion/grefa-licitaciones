"""Exportación de borradores markdown a Word (.docx) y PDF con formato de pliego."""

from __future__ import annotations

import io
import re
from typing import Any

DEFAULT_FORMATO: dict[str, Any] = {
    "fuente": "Arial",
    "tamano": 11,
    "interlineado": 1.15,
    "margen_cm": 2.5,
}


def parse_formato_desde_exigencias(exigencias: str) -> dict[str, Any]:
    """Intenta leer fuente, tamaño y márgenes desde el texto de exigencias."""
    fmt = dict(DEFAULT_FORMATO)
    if not exigencias:
        return fmt
    texto = exigencias

    m_fuente = re.search(
        r"(?:fuente|tipograf[ií]a|font)\s*[:\-]?\s*"
        r"(Arial|Times New Roman|Calibri|Helvetica|Courier New|Verdana|Georgia)",
        texto,
        flags=re.IGNORECASE,
    )
    if m_fuente:
        fmt["fuente"] = m_fuente.group(1).title().replace("New Roman", "New Roman")
        if "times" in m_fuente.group(1).lower():
            fmt["fuente"] = "Times New Roman"

    m_tam = re.search(
        r"(?:tama[nñ]o|size|cuerpo)\s*(?:de\s*(?:fuente|letra))?\s*[:\-]?\s*(\d{1,2})\s*(?:pt|puntos)?",
        texto,
        flags=re.IGNORECASE,
    )
    if m_tam:
        fmt["tamano"] = max(8, min(16, int(m_tam.group(1))))

    m_margen = re.search(
        r"m[aá]rgen(?:es)?\s*(?:de\s*)?[:\-]?\s*(\d+(?:[.,]\d+)?)\s*cm",
        texto,
        flags=re.IGNORECASE,
    )
    if m_margen:
        fmt["margen_cm"] = float(m_margen.group(1).replace(",", "."))

    m_inter = re.search(
        r"interlineado\s*[:\-]?\s*(\d+(?:[.,]\d+)?)",
        texto,
        flags=re.IGNORECASE,
    )
    if m_inter:
        fmt["interlineado"] = float(m_inter.group(1).replace(",", "."))

    return fmt


def _lineas_markdown(texto: str) -> list[tuple[str, str]]:
    """Devuelve lista (tipo, contenido): h1/h2/h3/li/p."""
    salida: list[tuple[str, str]] = []
    for cruda in (texto or "").splitlines():
        linea = cruda.rstrip()
        if not linea.strip():
            continue
        if linea.startswith("### "):
            salida.append(("h3", linea[4:].strip()))
        elif linea.startswith("## "):
            salida.append(("h2", linea[3:].strip()))
        elif linea.startswith("# "):
            salida.append(("h1", linea[2:].strip()))
        elif re.match(r"^\s*[-*]\s+", linea):
            salida.append(("li", re.sub(r"^\s*[-*]\s+", "", linea).strip()))
        elif re.match(r"^\s*\d+\.\s+", linea):
            salida.append(("li", re.sub(r"^\s*\d+\.\s+", "", linea).strip()))
        else:
            salida.append(("p", linea.strip()))
    return salida


def markdown_a_docx(
    markdown: str,
    *,
    titulo: str = "Borrador GREFA",
    formato: dict[str, Any] | None = None,
) -> bytes:
    try:
        from docx import Document
        from docx.enum.text import WD_LINE_SPACING
        from docx.shared import Cm, Pt
    except ImportError as exc:
        raise RuntimeError(
            "Falta python-docx. Instálalo con: pip install python-docx"
        ) from exc

    fmt = {**DEFAULT_FORMATO, **(formato or {})}
    doc = Document()
    for section in doc.sections:
        m = float(fmt.get("margen_cm") or 2.5)
        section.top_margin = Cm(m)
        section.bottom_margin = Cm(m)
        section.left_margin = Cm(m)
        section.right_margin = Cm(m)

    fuente = str(fmt.get("fuente") or "Arial")
    tamano = float(fmt.get("tamano") or 11)
    inter = float(fmt.get("interlineado") or 1.15)

    estilo = doc.styles["Normal"]
    estilo.font.name = fuente
    estilo.font.size = Pt(tamano)
    estilo.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    estilo.paragraph_format.line_spacing = inter

    h = doc.add_heading(titulo or "Borrador GREFA", level=0)
    for run in h.runs:
        run.font.name = fuente

    for tipo, contenido in _lineas_markdown(markdown):
        if tipo == "h1":
            p = doc.add_heading(contenido, level=1)
        elif tipo == "h2":
            p = doc.add_heading(contenido, level=2)
        elif tipo == "h3":
            p = doc.add_heading(contenido, level=3)
        elif tipo == "li":
            p = doc.add_paragraph(contenido, style="List Bullet")
        else:
            p = doc.add_paragraph(contenido)
        for run in p.runs:
            run.font.name = fuente
            if tipo == "p" or tipo == "li":
                run.font.size = Pt(tamano)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def markdown_a_pdf(
    markdown: str,
    *,
    titulo: str = "Borrador GREFA",
    formato: dict[str, Any] | None = None,
) -> bytes:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import cm
        from reportlab.platypus import ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer
    except ImportError as exc:
        raise RuntimeError(
            "Falta reportlab. Instálalo con: pip install reportlab"
        ) from exc

    fmt = {**DEFAULT_FORMATO, **(formato or {})}
    margen = float(fmt.get("margen_cm") or 2.5) * cm
    fuente = str(fmt.get("fuente") or "Helvetica")
    # reportlab built-in fonts
    mapa = {
        "arial": "Helvetica",
        "helvetica": "Helvetica",
        "times new roman": "Times-Roman",
        "times": "Times-Roman",
        "courier new": "Courier",
        "courier": "Courier",
    }
    font_name = mapa.get(fuente.lower(), "Helvetica")
    tamano = float(fmt.get("tamano") or 11)
    leading = tamano * float(fmt.get("interlineado") or 1.15)

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=margen,
        rightMargin=margen,
        topMargin=margen,
        bottomMargin=margen,
        title=titulo or "Borrador GREFA",
    )
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="CuerpoGREFA",
            fontName=font_name,
            fontSize=tamano,
            leading=leading,
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="H1GREFA",
            fontName=font_name,
            fontSize=tamano + 4,
            leading=leading + 4,
            spaceBefore=10,
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="H2GREFA",
            fontName=font_name,
            fontSize=tamano + 2,
            leading=leading + 2,
            spaceBefore=8,
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="H3GREFA",
            fontName=font_name,
            fontSize=tamano + 1,
            leading=leading + 1,
            spaceBefore=6,
            spaceAfter=4,
        )
    )

    def _esc(t: str) -> str:
        return (
            t.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
        )

    story: list[Any] = [
        Paragraph(_esc(titulo or "Borrador GREFA"), styles["H1GREFA"]),
        Spacer(1, 0.3 * cm),
    ]
    bullets: list[Any] = []

    def flush_bullets() -> None:
        nonlocal bullets
        if bullets:
            story.append(
                ListFlowable(
                    bullets,
                    bulletType="bullet",
                    start="•",
                )
            )
            bullets = []

    for tipo, contenido in _lineas_markdown(markdown):
        if tipo == "li":
            bullets.append(
                ListItem(Paragraph(_esc(contenido), styles["CuerpoGREFA"]))
            )
            continue
        flush_bullets()
        if tipo == "h1":
            story.append(Paragraph(_esc(contenido), styles["H1GREFA"]))
        elif tipo == "h2":
            story.append(Paragraph(_esc(contenido), styles["H2GREFA"]))
        elif tipo == "h3":
            story.append(Paragraph(_esc(contenido), styles["H3GREFA"]))
        else:
            story.append(Paragraph(_esc(contenido), styles["CuerpoGREFA"]))
    flush_bullets()

    doc.build(story)
    return buf.getvalue()


def construir_paquete_markdown(
    *,
    expediente: str,
    titulo: str,
    bloques: dict[str, str],
) -> str:
    """Une borradores admin/eco/tec en un único markdown."""
    partes = [
        f"# Paquete de presentación GREFA",
        f"",
        f"- **Expediente:** {expediente or '—'}",
        f"- **Objeto:** {titulo or '—'}",
        f"",
        (
            "> Borrador unificado (Administrativo + Económico + Técnico). "
            "Revisar modelos oficiales del pliego antes de presentar."
        ),
        "",
    ]
    etiquetas = {
        "admin": "Documentación administrativa",
        "eco": "Oferta económica",
        "tec": "Oferta técnica",
    }
    for clave in ("admin", "eco", "tec"):
        texto = (bloques.get(clave) or "").strip()
        partes.append(f"# {etiquetas[clave]}")
        partes.append("")
        partes.append(texto if texto else "_Sin borrador generado para este bloque._")
        partes.append("")
    return "\n".join(partes)
